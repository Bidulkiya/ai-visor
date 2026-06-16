"""범용 문서 파서 — 타입별로 텍스트를 추출해 공통 "구획(section)" 구조로 정규화한다.

PPTX 전용이던 추출을 PDF·DOCX·TXT·MD까지 일반화한다(HWP 제외). 모든 타입을
slides.ts의 Slide와 호환되는 dict로 내보낸다:
  {"number","title","bodyText","speakerNotes"}  (imageDataUrl은 server가 채운다)

설계 원칙(slide_renderer.py와 같은 정신):
- 선택 의존(PyMuPDF·python-docx)이 없으면 **크래시하지 않고** UnsupportedDocument로
  사유를 알린다 — 호출자(server)는 사용자에게 명확히 안내한다.
- 지원하지 않는 확장자(HWP 등)도 UnsupportedDocument로 명확히 알린다(크래시 금지).
- 대용량 문서는 구획 수·총 문자 수를 상한으로 자른다(메모리·IPC 폭주 방지) — 잘릴 땐
  마지막에 안내 구획을 덧붙인다.
"""
from __future__ import annotations

import os

# 파서가 다룰 수 있는 확장자 → 문서 타입. 여기 없는 확장자(.hwp 등)는 지원 안 함.
_EXTENSION_TO_TYPE = {
    ".pptx": "pptx",
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
}

# 대용량 방어 상한 — 발표·이해 모두에 충분하되 폭주는 막는다
MAX_SECTIONS = 300
MAX_SECTION_CHARS = 8000
MAX_TOTAL_CHARS = 600_000
# 구획 제목으로 쓸 첫 줄의 최대 길이
_TITLE_MAX_CHARS = 80
# TXT를 구획으로 나눌 때의 목표 크기(문단 경계에서 묶는다)
_TXT_CHUNK_TARGET_CHARS = 1500


class UnsupportedDocument(Exception):
    """지원하지 않는 형식이거나 필요한 파서 의존성이 없을 때. server가 사유를 사용자에게 전달."""


def detect_document_type(path: str):
    """확장자로 문서 타입을 판별한다. 지원 안 하면 None."""
    return _EXTENSION_TO_TYPE.get(os.path.splitext(path)[1].lower())


def _first_line_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if stripped:
            return stripped[:_TITLE_MAX_CHARS]
    return fallback


def _read_text_file(path: str) -> str:
    """한글 문서까지 견디게 인코딩을 차례로 시도한다(utf-8 → cp949 → latin-1 대체)."""
    for encoding in ("utf-8", "cp949"):
        try:
            with open(path, "r", encoding=encoding) as handle:
                return handle.read()
        except (UnicodeDecodeError, LookupError):
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read()


def _section(number: int, title: str, body_text: str) -> dict:
    return {
        "number": number,
        "title": title.strip()[:_TITLE_MAX_CHARS],
        "bodyText": body_text.strip()[:MAX_SECTION_CHARS],
        "speakerNotes": "",
    }


# ── 타입별 추출 ── (각자 [{number,title,bodyText,speakerNotes}] 리스트를 반환)

def _extract_pdf(path: str) -> list:
    try:
        import fitz  # PyMuPDF
    except ImportError as error:
        raise UnsupportedDocument(
            "PDF를 읽으려면 PyMuPDF가 필요합니다. 'pip install pymupdf' 후 다시 시도해 주세요."
        ) from error
    try:
        # 확장자만 .pdf이고 내용이 손상·암호화·다른 형식이면 여기서 실패 — 미지원으로 명확히 알린다
        document = fitz.open(path)
    except (fitz.FileDataError, RuntimeError, ValueError) as error:
        raise UnsupportedDocument("PDF를 열 수 없어요 — 손상됐거나 형식이 올바르지 않아요.") from error
    try:
        sections = []
        for index in range(document.page_count):
            page_text = document.load_page(index).get_text().strip()
            number = index + 1
            title = _first_line_title(page_text, "페이지 %d" % number)
            sections.append(_section(number, title, page_text))
        return sections
    finally:
        document.close()


def _docx_table_text(table) -> str:
    lines = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        joined = " | ".join(cell for cell in cells if cell)
        if joined:
            lines.append(joined)
    return "\n".join(lines)


def _extract_docx(path: str) -> list:
    try:
        import docx  # python-docx
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as error:
        raise UnsupportedDocument(
            "DOCX를 읽으려면 python-docx가 필요합니다. 'pip install python-docx' 후 다시 시도해 주세요."
        ) from error
    try:
        # 확장자만 .docx이고 내용이 손상·다른 형식이면 PackageNotFoundError — 미지원으로 명확히
        document = docx.Document(path)
    except PackageNotFoundError as error:
        raise UnsupportedDocument("Word 문서를 열 수 없어요 — 손상됐거나 형식이 올바르지 않아요.") from error
    # 제목 스타일(Heading)마다 새 구획을 연다. 제목 전 본문은 첫 구획에 모은다.
    sections = []
    current_title = ""
    current_body = []

    def flush() -> None:
        if current_title or current_body:
            number = len(sections) + 1
            title = current_title or "본문"
            sections.append(_section(number, title, "\n".join(current_body)))

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
        if (style_name.startswith("Heading") or style_name == "Title") and text:
            flush()
            current_title = text
            current_body = []
        elif text:
            current_body.append(text)
    flush()

    # 표는 본문 뒤에 별도 구획으로 덧붙인다(셀 텍스트 보존)
    for table in document.tables:
        table_text = _docx_table_text(table)
        if table_text:
            sections.append(_section(len(sections) + 1, "표", table_text))

    if not sections:
        sections.append(_section(1, "본문", ""))
    return sections


def _chunk_paragraphs(text: str) -> list:
    """빈 줄로 나뉜 문단을 목표 크기로 묶어 구획 리스트로 만든다(TXT용)."""
    paragraphs = [block.strip() for block in text.split("\n\n") if block.strip()]
    if not paragraphs:
        return []
    chunks = []
    buffer = []
    buffer_len = 0
    for paragraph in paragraphs:
        if buffer and buffer_len + len(paragraph) > _TXT_CHUNK_TARGET_CHARS:
            chunks.append("\n\n".join(buffer))
            buffer = []
            buffer_len = 0
        buffer.append(paragraph)
        buffer_len += len(paragraph)
    if buffer:
        chunks.append("\n\n".join(buffer))
    return chunks


def _extract_markdown(text: str) -> list:
    """'#' 제목을 경계로 구획을 나눈다(MD 구조 인식). 제목 전 내용은 '본문' 구획."""
    sections = []
    current_title = ""
    current_body = []

    def flush() -> None:
        if current_title or any(line.strip() for line in current_body):
            number = len(sections) + 1
            sections.append(_section(number, current_title or "본문", "\n".join(current_body)))

    for line in text.splitlines():
        stripped = line.lstrip()
        if stripped.startswith("#"):
            heading = stripped.lstrip("#").strip()
            if heading:
                flush()
                current_title = heading
                current_body = []
                continue
        current_body.append(line)
    flush()
    return sections


def _extract_text(path: str, doc_type: str) -> list:
    text = _read_text_file(path)
    base_name = os.path.splitext(os.path.basename(path))[0]
    if doc_type == "md":
        sections = _extract_markdown(text)
        if sections:
            return sections
    # TXT(또는 제목 없는 MD): 문단을 크기로 묶어 구획화
    chunks = _chunk_paragraphs(text)
    if not chunks:
        return [_section(1, base_name or "본문", text)]
    if len(chunks) == 1:
        return [_section(1, _first_line_title(chunks[0], base_name or "본문"), chunks[0])]
    return [_section(index + 1, "구획 %d" % (index + 1), chunk) for index, chunk in enumerate(chunks)]


def _apply_limits(sections: list) -> list:
    """구획 수·총 문자 수 상한을 적용한다. 잘리면 안내 구획을 덧붙인다(조용한 손실 방지)."""
    truncated = False
    limited = sections
    if len(limited) > MAX_SECTIONS:
        limited = limited[:MAX_SECTIONS]
        truncated = True
    total = 0
    result = []
    for section in limited:
        body = section.get("bodyText", "")
        if total + len(body) > MAX_TOTAL_CHARS:
            allowed = max(0, MAX_TOTAL_CHARS - total)
            section = dict(section, bodyText=body[:allowed])
            truncated = True
            result.append(section)
            break
        total += len(body)
        result.append(section)
    if truncated:
        result.append(
            {
                "number": len(result) + 1,
                "title": "(이후 생략)",
                "bodyText": "문서가 너무 길어 일부만 읽었어요. 핵심 위주로 봐 주세요.",
                "speakerNotes": "",
            }
        )
    # 번호를 1부터 다시 매겨 연속성을 보장한다
    for index, section in enumerate(result):
        section["number"] = index + 1
    return result


def extract_document(path: str) -> dict:
    """문서를 타입에 맞게 파싱해 공통 구조로 돌려준다.

    반환: {"docType": str, "sections": [{number,title,bodyText,speakerNotes}, ...]}
    지원 안 하는 형식·의존성 누락은 UnsupportedDocument를 던진다(server가 안내).
    """
    doc_type = detect_document_type(path)
    if doc_type is None:
        extension = os.path.splitext(path)[1].lower() or "(없음)"
        raise UnsupportedDocument(
            "지원하지 않는 형식입니다(%s). PDF·DOCX·PPTX·TXT·MD만 열 수 있어요." % extension
        )
    if doc_type == "pptx":
        from pptx_parser import extract_slides

        sections = extract_slides(path)
    elif doc_type == "pdf":
        sections = _extract_pdf(path)
    elif doc_type == "docx":
        sections = _extract_docx(path)
    else:
        sections = _extract_text(path, doc_type)
    return {"docType": doc_type, "sections": _apply_limits(sections)}
